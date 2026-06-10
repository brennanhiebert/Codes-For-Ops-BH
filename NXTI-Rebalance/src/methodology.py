"""
methodology.py — generate a visual methodology document explaining:
  1. how target SHARE SIZE is determined  (share_sizing.png)
  2. the full BUY / SELL / IN-KIND / PARTIAL decision tree  (decision_tree.png)

and assemble both into a Word document with explanatory text.

Run:  python -m src.methodology
Outputs:
  output/diagrams/share_sizing.png
  output/diagrams/decision_tree.png
  output/NXTI_Methodology_<date>.docx
"""
from __future__ import annotations

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

from . import config

# ---- palette (matches the blotter workbook) ----
BUY = "#D9E1F2"; MARKET = "#C6EFCE"; INKIND = "#EDEDED"; AMBER = "#FFEB9C"
DECISION = "#FFFFFF"; NEUTRAL = "#F2F2F2"; NAVY = "#1F3864"; BLUE = "#2E5496"
EDGE = "#7F7F7F"; TXT = "#1F1F1F"

DIAG_DIR = config.OUTPUT_DIR / "diagrams"
DIAG_DIR.mkdir(exist_ok=True)


# --------------------------------------------------------------------------- #
# drawing helpers
# --------------------------------------------------------------------------- #
def _box(ax, x, y, w, h, text, fc=DECISION, ec=EDGE, fs=10, bold=False,
         style="round", text_color=TXT):
    p = FancyBboxPatch(
        (x - w / 2, y - h / 2), w, h,
        boxstyle=f"{style},pad=0.02,rounding_size=1.2"
        if style == "round" else f"{style},pad=0.02",
        linewidth=1.4, edgecolor=ec, facecolor=fc, zorder=2,
    )
    ax.add_patch(p)
    ax.text(x, y, text, ha="center", va="center", fontsize=fs,
            color=text_color, zorder=3,
            fontweight="bold" if bold else "normal", wrap=True)
    return (x, y, w, h)


def _arrow(ax, p_from, p_to, label=None, label_dx=0, label_dy=0,
           color=EDGE, side_from="bottom", side_to="top"):
    fx, fy, fw, fh = p_from
    tx, ty, tw, th = p_to
    def anchor(b, side):
        x, y, w, h = b
        return {
            "bottom": (x, y - h / 2), "top": (x, y + h / 2),
            "left": (x - w / 2, y), "right": (x + w / 2, y),
        }[side]
    a = anchor(p_from, side_from)
    b = anchor(p_to, side_to)
    arr = FancyArrowPatch(a, b, arrowstyle="-|>", mutation_scale=14,
                          linewidth=1.4, color=color, zorder=1,
                          connectionstyle="arc3,rad=0")
    ax.add_patch(arr)
    if label:
        mx, my = (a[0] + b[0]) / 2 + label_dx, (a[1] + b[1]) / 2 + label_dy
        ax.text(mx, my, label, ha="center", va="center", fontsize=8.5,
                style="italic", color="#404040",
                bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="none"),
                zorder=4)


def _canvas(w_in, h_in, title):
    fig, ax = plt.subplots(figsize=(w_in, h_in))
    ax.set_xlim(0, 100); ax.set_ylim(0, 100); ax.axis("off")
    ax.text(50, 97.5, title, ha="center", va="center", fontsize=16,
            fontweight="bold", color=NAVY)
    return fig, ax


# --------------------------------------------------------------------------- #
# Diagram 1 — share sizing
# --------------------------------------------------------------------------- #
def make_share_sizing(example: dict) -> str:
    fig, ax = _canvas(12, 7.6, "How Target Share Size Is Determined")

    # The formula band.
    _box(ax, 50, 87,  92, 8,
         "target_shares  =  index_weight  ×  investable_base  ÷  current_price",
         fc=BLUE, ec=NAVY, fs=14, bold=True, text_color="white")

    # Three inputs.
    iw = _box(ax, 18, 70, 28, 13,
              "index_weight  (precise)\n= index_shares × closing_price ÷ index_close\n"
              "the displayed 0.18% column is too coarse;\nwe rebuild full precision",
              fc=NEUTRAL, fs=9)
    ib = _box(ax, 50, 70, 28, 13,
              "investable_base\n= Σ NXTI equity Market Value\n"
              f"= ${example['investable_base']:,.0f}\n(cash buffer preserved)",
              fc=NEUTRAL, fs=9)
    cp = _box(ax, 82, 70, 28, 13,
              "current_price\n= live Bloomberg PX_LAST\n"
              "(falls back to BNY / proforma\nif the terminal is offline)",
              fc=NEUTRAL, fs=9)

    band = (50, 87, 92, 8)
    for b in (iw, ib, cp):
        _arrow(ax, band, b, side_from="bottom", side_to="top")

    # Pipeline.
    s1 = _box(ax, 50, 50, 60, 7,
              "Round to whole shares  →  TARGET shares for the name", fc="#FCE4D6", fs=10)
    for b in (iw, ib, cp):
        _arrow(ax, b, s1, side_from="bottom", side_to="top")

    s2 = _box(ax, 50, 39, 70, 7,
              "delta = target − current   (current = Quantity held in the tracker)",
              fc="#FCE4D6", fs=10)
    _arrow(ax, s1, s2)

    # Worked example panel.
    ex = example
    txt = (
        "WORKED EXAMPLE — AFLAC INC  (CUSIP 001055102)\n"
        f"index_shares = {ex['index_shares']:.6f}    closing = ${ex['closing']:.2f}    "
        f"index_close = {ex['index_close']:.2f}\n"
        f"precise weight = {ex['index_shares']:.6f} × {ex['closing']:.2f} ÷ "
        f"{ex['index_close']:.2f}  =  {ex['weight']:.5%}\n"
        f"target = {ex['weight']:.5%} × ${ex['investable_base']:,.0f} ÷ "
        f"${ex['price']:.2f}  =  {ex['target_raw']:.1f}  →  {ex['target']} shares\n"
        f"current held = {ex['current']}   →   delta = {ex['target']} − {ex['current']} "
        f"= {ex['delta']:+d}   →   {ex['action']}"
    )
    _box(ax, 50, 20, 92, 16, txt, fc="#FFF6DD", ec="#BF9000", fs=10.5)

    fig.tight_layout()
    path = DIAG_DIR / "share_sizing.png"
    fig.savefig(path, dpi=170, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return str(path)


# --------------------------------------------------------------------------- #
# Diagram 2 — decision tree
# --------------------------------------------------------------------------- #
def make_decision_tree() -> str:
    fig, ax = _canvas(14, 12, "Trade-Method Decision Tree  (per security, joined on CUSIP)")

    root = _box(ax, 50, 92, 40, 6,
                "Each security in pro-forma ∪ portfolio", fc=NAVY, fs=11,
                bold=True, text_color="white")

    # Membership branches.
    pf_only = _box(ax, 16, 80, 24, 6, "In pro-forma\nbut NOT held", fc=NEUTRAL, fs=9)
    both    = _box(ax, 50, 80, 24, 6, "In BOTH", fc=NEUTRAL, fs=9)
    tk_only = _box(ax, 84, 80, 24, 6, "Held but NOT\nin pro-forma", fc=NEUTRAL, fs=9)
    _arrow(ax, root, pf_only, side_to="top"); _arrow(ax, root, both)
    _arrow(ax, root, tk_only, side_to="top")

    # Classification.
    add  = _box(ax, 16, 70, 20, 5.5, "ADD", fc=BUY, fs=11, bold=True)
    cont = _box(ax, 50, 70, 20, 5.5, "CONTINUING", fc="#DDEBF7", fs=11, bold=True)
    drop = _box(ax, 84, 70, 20, 5.5, "DROP", fc=AMBER, fs=11, bold=True)
    _arrow(ax, pf_only, add); _arrow(ax, both, cont); _arrow(ax, tk_only, drop)

    # ADD terminal.
    add_buy = _box(ax, 16, 58, 24, 7,
                   "BUY full target\n@ MARKET", fc=BUY, fs=10, bold=True)
    _arrow(ax, add, add_buy)

    # CONTINUING comparison.
    cmp = _box(ax, 50, 59, 26, 6, "compare target vs current", fc=DECISION, fs=10)
    _arrow(ax, cont, cmp)
    c_buy = _box(ax, 36, 48, 19, 7, "target > current\nBUY delta @ MARKET",
                 fc=BUY, fs=9, bold=True)
    c_no  = _box(ax, 50, 48, 13, 7, "target =\ncurrent\nNO TRADE", fc=NEUTRAL, fs=9)
    c_par = _box(ax, 65, 48, 19, 7, "target < current\nPARTIAL SELL", fc=AMBER, fs=9,
                 bold=True)
    _arrow(ax, cmp, c_buy, side_to="top"); _arrow(ax, cmp, c_no)
    _arrow(ax, cmp, c_par, side_to="top")

    # DROP -> full sell.
    full = _box(ax, 84, 58, 24, 7, "FULL SELL\n(entire position)", fc=AMBER, fs=10,
                bold=True)
    _arrow(ax, drop, full)

    # Funnel sells into lot engine.
    lot = _box(ax, 74, 38, 40, 7,
               "LOT ENGINE — examine EACH tax lot\n(buys never reach here)",
               fc=NAVY, fs=10, bold=True, text_color="white")
    _arrow(ax, c_par, lot, side_from="bottom", side_to="left")
    _arrow(ax, full, lot, side_from="bottom", side_to="top")

    # Lot-level decision.
    q = _box(ax, 74, 27, 34, 6, "current_price  <  lot's cost basis ?", fc=DECISION, fs=10)
    _arrow(ax, lot, q)

    mkt = _box(ax, 55, 14, 28, 9,
               "LOSS  →  SELL AT MARKET\nrealize the loss\norder: HIGHEST basis first",
               fc=MARKET, fs=9.5, bold=True)
    ink = _box(ax, 90, 14, 28, 9,
               "GAIN  →  TRANSFER IN-KIND\navoid realizing the gain\norder: LOWEST basis first",
               fc=INKIND, fs=9.5, bold=True)
    _arrow(ax, q, mkt, label="YES (loss)", side_from="bottom", side_to="top",
           label_dx=-4)
    _arrow(ax, q, ink, label="NO (gain)", side_from="bottom", side_to="top",
           label_dx=4)

    # Notes on full vs partial.
    _box(ax, 24, 30, 34, 13,
         "FULL SELL (DROP)\n• every lot is placed\n• losses → market leg\n"
         "• gains → in-kind leg\n• up to TWO legs per name",
         fc="#FBE5D6", ec="#C55A11", fs=9)
    _box(ax, 24, 13, 34, 13,
         "PARTIAL SELL (trim)\n• cover exactly |delta|, then stop\n"
         f"• default: drain LOSS pool first\n  ({config.PARTIAL_SELL_PRIORITY}),"
         " then gains\n• split the last lot if needed",
         fc="#FBE5D6", ec="#C55A11", fs=9)

    fig.tight_layout()
    path = DIAG_DIR / "decision_tree.png"
    fig.savefig(path, dpi=170, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return str(path)


# --------------------------------------------------------------------------- #
# worked-example numbers (computed from the files; no Bloomberg needed)
# --------------------------------------------------------------------------- #
def _example() -> dict:
    from . import loaders
    pf, meta = loaders.load_proforma()
    tk, tkm = loaders.load_tracker()
    cusip = "001055102"  # AFLAC
    pr = pf.set_index("cusip").loc[cusip]
    cur = float(tk.set_index("cusip").loc[cusip, "quantity"])
    base = tkm["investable_base"]
    weight = float(pr["index_weight"])
    closing = float(pr["closing_price"])
    price = closing  # use closing for a deterministic, reproducible example
    target_raw = weight * base / price
    target = round(target_raw)
    return {
        "investable_base": base, "index_shares": float(pr["index_shares"]),
        "closing": closing, "index_close": meta["index_close"], "weight": weight,
        "price": price, "target_raw": target_raw, "target": target,
        "current": int(cur), "delta": int(target - cur),
        "action": "BUY" if target > cur else "SELL" if target < cur else "NO TRADE",
    }


# --------------------------------------------------------------------------- #
# Word document
# --------------------------------------------------------------------------- #
def build_doc() -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    print("[methodology] computing worked example from input files ...")
    ex = _example()
    print("[methodology] rendering diagrams ...")
    sizing_png = make_share_sizing(ex)
    tree_png = make_decision_tree()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(11)

    def heading(txt, size=16, color=RGBColor(0x1F, 0x38, 0x64)):
        p = doc.add_paragraph()
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
        return p

    # Title page-ish header.
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("NXTI Rebalance — Trade-Method Methodology")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(f"How the engine sizes positions and routes every trade  ·  "
                     f"run date {config.RUN_DATE}")
    rs.italic = True; rs.font.size = Pt(11); rs.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    doc.add_paragraph()

    intro = doc.add_paragraph()
    intro.add_run(
        "This document explains the two decisions the engine makes: (1) HOW MANY shares "
        "each name should hold after the rebalance, and (2) for every resulting trade, "
        "WHICH METHOD to use — buy, sell at market, or transfer in-kind — decided lot by "
        "lot. The tool only produces a blotter for human review; it never places a trade."
    )

    # Section 1.
    heading("1.  Determining share size")
    doc.add_paragraph(
        "The pro-forma's “Index Shares” column is an index-level figure (e.g. AFLAC ≈ "
        "0.048), not a portfolio share count (AFLAC is actually held = 614). Using it "
        "literally would liquidate ~99% of every continuing name. Instead each name is "
        "sized to the fund using its index WEIGHT, the fund's investable base, and the "
        "live price:"
    )
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run("target_shares = index_weight × investable_base ÷ current_price")
    rr.bold = True; rr.font.size = Pt(12)
    doc.add_picture(sizing_png, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Note the precision fix: the displayed “Index Weighting” column is rounded to two "
        "decimals (0.18%), too coarse for sizing, so the full-precision weight is rebuilt "
        "from index_shares × closing_price ÷ index_close (these recovered weights sum to "
        "exactly 100%). The investable base defaults to the sum of NXTI equity market "
        "value, preserving the cash buffer."
    )

    # Section 2.
    heading("2.  Determining the trade method (buy / sell / in-kind)")
    doc.add_paragraph(
        "Membership and the share delta classify each name and derive its action. Buys are "
        "always at market. Sells — whether a full exit (DROP) or a partial trim — go "
        "through the LOT ENGINE, which inspects each tax lot individually:"
    )
    b1 = doc.add_paragraph(style="List Bullet")
    b1.add_run("Lot at a loss ").bold = True
    b1.add_run("(current price < the lot's basis) → SELL AT MARKET to realize the loss; "
               "loss lots are ordered highest-basis-first (biggest loss per share first).")
    b2 = doc.add_paragraph(style="List Bullet")
    b2.add_run("Lot at a gain ").bold = True
    b2.add_run("(current price ≥ the lot's basis) → TRANSFER IN-KIND to avoid realizing "
               "the gain; gain lots are ordered lowest-basis-first (biggest embedded gain "
               "shielded first).")
    doc.add_paragraph(
        "A full sell places every lot, so one name can produce two legs (a market leg and "
        "an in-kind leg). A partial trim covers only |delta| shares, draining the loss "
        "pool first by default (configurable), then taking gains in-kind, splitting the "
        "final lot if it is only partially needed."
    )
    doc.add_picture(tree_png, width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    rc = cap.add_run("Both the flat blotter and the expandable blotter (with collapsible "
                     "per-lot detail) are produced from exactly this logic.")
    rc.italic = True; rc.font.size = Pt(9.5); rc.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    path = config.OUTPUT_DIR / f"NXTI_Methodology_{config.run_stamp()}.docx"
    try:
        doc.save(path)
    except PermissionError:
        import datetime as _dt
        path = path.with_name(f"{path.stem}_{_dt.datetime.now():%H%M%S}{path.suffix}")
        doc.save(path)
    print(f"[methodology] saved -> {path}")
    print(f"[methodology] diagrams -> {sizing_png}\n                          {tree_png}")
    return str(path)


if __name__ == "__main__":
    build_doc()
